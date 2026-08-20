from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


PATH = Path(r"C:\Users\sung4\OneDrive\바탕 화면\EH\ESA\output\manual\ESVA_공개교보재_도면분석_사용자매뉴얼_v0.2.2.docx")
required = [
    "공개 교보재 기반\n사용자 매뉴얼",
    "Wikimedia 단선도·배선도",
    "83쪽 공공기관 설계도면",
    "검색과 AI 답변",
    "도면 분석: 판독 구조",
    "Pxx-Axx",
    "Pxx-Cxxx",
    "Pxx-Uxxx",
    "Google 모델 선택 가이드",
    "작업 완료 체크리스트",
]
forbidden = ["TODO", "TBD", "Lorem ipsum", "페이지를 찾을 수 없습니다"]

with ZipFile(PATH) as zf:
    bad = zf.testzip()
    assert bad is None, f"corrupt member: {bad}"
    media = [name for name in zf.namelist() if name.startswith("word/media/")]
    assert len(media) == 6, f"expected 6 media items, got {len(media)}"
    assert all(len(zf.read(name)) > 1000 for name in media), "empty image payload"
    xml = zf.read("word/document.xml").decode("utf-8")
    page_breaks = xml.count('w:type="page"')
    assert page_breaks >= 12, f"expected at least 12 page breaks, got {page_breaks}"
    alt_count = xml.count("descr=")
    assert alt_count >= 6, f"expected 6 alt descriptions, got {alt_count}"

doc = Document(PATH)
text = "\n".join(p.text for p in doc.paragraphs)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            text += "\n" + "\n".join(p.text for p in cell.paragraphs)

for token in required:
    assert token in text, f"missing required token: {token}"
for token in forbidden:
    assert token not in text, f"forbidden token remains: {token}"

headings = [p.text for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")]
captions = [p.text for p in doc.paragraphs if p.style and p.style.name == "Caption"]
assert len(doc.inline_shapes) == 6, f"inline shapes: {len(doc.inline_shapes)}"
assert len(captions) == 6, f"captions: {len(captions)}"
assert len(doc.tables) >= 15, f"tables: {len(doc.tables)}"
assert len(headings) >= 20, f"headings: {len(headings)}"
assert len(doc.sections) == 1, f"sections: {len(doc.sections)}"

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            tcw = cell._tc.tcPr.find(qn("w:tcW")) if cell._tc.tcPr is not None else None
            assert tcw is not None, "table cell width missing"

print(f"PASS file={PATH.name}")
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} headings={len(headings)}")
print(f"images={len(doc.inline_shapes)} captions={len(captions)} page_breaks={page_breaks} alt={alt_count}")
