from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]  # scripts/manual/ 기준 리포 루트 — 절대경로 하드코딩은 이 PC 밖에서 죽는다
OUT = ROOT / "output" / "manual" / "ESVA_공개교보재_도면분석_사용자매뉴얼_v0.2.2.docx"
TMP = ROOT / "tmp" / "manual"
FIG = ROOT / "fixtures" / "drawings" / "external" / "wiki-oneline.png"

BLUE = "1F4D78"
BLUE2 = "2E74B5"
NAVY = "163A63"
ORANGE = "D97706"
GREEN = "178A5B"
RED = "B42318"
INK = "262626"
MUTED = "666666"
PALE = "F4F7FA"
PALE_ORANGE = "FFF7ED"
PALE_GREEN = "EFFAF5"
PALE_RED = "FEF3F2"
RULE = "D7DEE7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def set_run_font(run, latin="Calibri", korean="Malgun Gothic") -> None:
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), korean)


def set_paragraph_text(p, text: str, *, bold=False, color=INK, size=None) -> None:
    p.clear()
    r = p.add_run(text)
    r.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    if size:
        r.font.size = Pt(size)
    set_run_font(r)


def shade_paragraph(p, fill: str, border: str | None = None) -> None:
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border:
        pbdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border)
        pbdr.append(left)
        p_pr.append(pbdr)


def add_note(doc, title: str, body: str, *, kind="info"):
    colors = {
        "info": (PALE, BLUE2),
        "tip": (PALE_GREEN, GREEN),
        "warn": (PALE_ORANGE, ORANGE),
        "danger": (PALE_RED, RED),
    }
    fill, accent = colors[kind]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    table.style = "Table Grid"
    repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(accent)
    set_run_font(r)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.18
    for r2 in p2.runs:
        set_run_font(r2)
        r2.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def set_alt_text(shape, descr: str) -> None:
    inline = shape._inline
    doc_pr = inline.docPr
    doc_pr.set("descr", descr)


def add_figure(doc, image_path: Path, caption: str, alt: str, width=6.15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(image_path), width=Inches(width))
    set_alt_text(shape, alt)
    cp = doc.add_paragraph(style="Caption")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.keep_together = True
    cp.paragraph_format.space_after = Pt(8)
    r = cp.add_run(caption)
    r.italic = True
    r.font.color.rgb = RGBColor.from_string(MUTED)
    r.font.size = Pt(9)
    set_run_font(r)


def add_bullets(doc, items: list[str], style="List Bullet"):
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)
        for r in p.runs:
            set_run_font(r)


def add_steps(doc, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(5)
        p.add_run(item)
        for r in p.runs:
            set_run_font(r)


def add_section_title(doc, number: str, title: str, subtitle: str | None = None):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    r1 = p.add_run(number + "  ")
    r1.bold = True
    r1.font.color.rgb = RGBColor.from_string(ORANGE)
    r1.font.size = Pt(13)
    set_run_font(r1)
    r2 = p.add_run(title)
    r2.bold = True
    r2.font.color.rgb = RGBColor.from_string(NAVY)
    r2.font.size = Pt(22)
    set_run_font(r2)
    if subtitle:
        sp = doc.add_paragraph(subtitle)
        sp.paragraph_format.space_after = Pt(12)
        for r in sp.runs:
            set_run_font(r)
            r.font.color.rgb = RGBColor.from_string(MUTED)


def add_page_break(doc):
    doc.add_page_break()


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int], header_fill="E8EEF5"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], header_fill)
        p = table.rows[0].cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_paragraph_text(p, header, bold=True, color=NAVY, size=9.5)
    for row_data in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for idx, value in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            set_paragraph_text(p, value, color=INK, size=9.2)
    return table


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_run_font(run)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    run.font.size = Pt(8)


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.72)
sec.bottom_margin = Inches(0.72)
sec.left_margin = Inches(1.0)
sec.right_margin = Inches(1.0)
sec.different_first_page_header_footer = True

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

for style_name, size, color, before, after in (
    ("Heading 1", 16, BLUE2, 18, 10),
    ("Heading 2", 13, BLUE2, 14, 7),
    ("Heading 3", 11.5, BLUE, 10, 5),
):
    st = styles[style_name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True
    st_rpr = st._element.get_or_add_rPr()
    st_rfonts = st_rpr.rFonts
    if st_rfonts is None:
        st_rfonts = OxmlElement("w:rFonts")
        st_rpr.insert(0, st_rfonts)
    st_rfonts.set(qn("w:eastAsia"), "Malgun Gothic")

if "Title Section" not in styles:
    ts = styles.add_style("Title Section", WD_STYLE_TYPE.PARAGRAPH)
else:
    ts = styles["Title Section"]
ts.paragraph_format.space_before = Pt(0)
ts.paragraph_format.space_after = Pt(6)
ts.paragraph_format.keep_with_next = True

for list_name in ("List Bullet", "List Number"):
    st = styles[list_name]
    st.font.size = Pt(10.5)
    st.paragraph_format.left_indent = Inches(0.375)
    st.paragraph_format.first_line_indent = Inches(-0.188)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.25
    st_rpr = st._element.get_or_add_rPr()
    st_rfonts = st_rpr.rFonts
    if st_rfonts is None:
        st_rfonts = OxmlElement("w:rFonts")
        st_rpr.insert(0, st_rfonts)
    st_rfonts.set(qn("w:eastAsia"), "Malgun Gothic")

caption = styles["Caption"]
caption.font.name = "Calibri"
caption_rpr = caption._element.get_or_add_rPr()
caption_rfonts = caption_rpr.rFonts
if caption_rfonts is None:
    caption_rfonts = OxmlElement("w:rFonts")
    caption_rpr.insert(0, caption_rfonts)
caption_rfonts.set(qn("w:eastAsia"), "Malgun Gothic")

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = header.add_run("ESVA 공개 교보재 기반 사용자 매뉴얼  |  v0.2.2")
hr.font.size = Pt(8)
hr.font.color.rgb = RGBColor.from_string(MUTED)
set_run_font(hr)
page_number(sec.footer.paragraphs[0])

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(82)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("⚡  ESVA")
r.bold = True
r.font.size = Pt(25)
r.font.color.rgb = RGBColor.from_string(NAVY)
set_run_font(r)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(22)
p.paragraph_format.space_after = Pt(10)
r = p.add_run("공개 교보재 기반\n사용자 매뉴얼")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = RGBColor.from_string(NAVY)
set_run_font(r)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(26)
r = p.add_run("단선도 · 분전반 · 전력/조명/제어 도면 · DXF 분석")
r.font.size = Pt(13)
r.font.color.rgb = RGBColor.from_string(BLUE2)
set_run_font(r)

table = doc.add_table(rows=4, cols=2)
set_table_geometry(table, [2500, 6860])
table.style = "Table Grid"
repeat_table_header(table.rows[0])
cover_rows = [
    ("대상", "초급~고급 전기 직무자, 설계·시공·안전·검토 담당자"),
    ("사용 교보재", "Wikimedia 단선도·배선도, 수변전/기능사 교재, 공공기관 전기설계 PDF, 합성 DXF L1~L3"),
    ("도면 수준", "교육용 단순 회로부터 분전반·전력·조명·제어 및 83쪽 공공기관 설계도면까지"),
    ("버전", "v0.2.2  |  2026-07-23"),
]
for idx, (label, value) in enumerate(cover_rows):
    set_cell_shading(table.cell(idx, 0), PALE)
    set_paragraph_text(table.cell(idx, 0).paragraphs[0], label, bold=True, color=NAVY, size=9.5)
    set_paragraph_text(table.cell(idx, 1).paragraphs[0], value, color=INK, size=9.5)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(42)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("핵심 원칙")
r.bold = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor.from_string(ORANGE)
set_run_font(r)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("대표 검증 범위: Wikipedia 단선도, 대산전기 11쪽, KIMM 18·83쪽, KBRI 9쪽, CODIL 466쪽, 박문각 18쪽, 합성 DXF L1~L3")
r.font.size = Pt(10.5)
r.font.color.rgb = RGBColor.from_string(MUTED)
set_run_font(r)

add_page_break(doc)

# Quick overview
add_section_title(doc, "00", "한눈에 보기", "처음 쓰는 사람은 이 두 페이지만 읽어도 기본 작업을 시작할 수 있습니다.")
add_figure(
    doc,
    TMP / "live-home.png",
    "그림 1. 로그인하지 않은 상태의 ESVA 홈: 검색창에서 질문·계산 조건·도면 검토 내용을 입력한다.",
    "ESVA 홈 화면. 상단 메뉴, 중앙 질문 입력창, 최근 스레드와 도구 바로가기가 보인다.",
    width=6.15,
)

doc.add_heading("기능 지도", level=2)
add_table(
    doc,
    ["메뉴", "주요 용도", "AI 키"],
    [
        ["검색", "전기 질문, 기준서 조항 탐색, 계산 의도 자동 분류", "답변 생성 시 필요"],
        ["계산기", "전압강하·케이블·변압기·보호협조 등 57개 수식 계산", "불필요"],
        ["현장안전", "작업 조건을 안전 체크리스트와 작업 전 확인 항목으로 변환", "필요"],
        ["Studio", "명판·문서·이미지에서 입력 후보를 추출하고 계산으로 연결", "필요"],
        ["SLD", "이미지·DXF·PDF 단선도 전체 판독, 관계·계산·제안", "이미지 정밀판독 시 필요"],
        ["기준서/용어", "KEC·NEC·IEC 스냅샷과 전기 용어 탐색", "불필요"],
    ],
    [1500, 6060, 1800],
)

add_note(
    doc,
    "5분 빠른 시작",
    "① 홈에서 질문을 입력하거나 메뉴를 선택합니다. ② AI가 필요한 작업이면 설정에서 개인 API 키(BYOK)를 등록합니다. ③ 계산 결과는 입력값·단위·판본을 확인하고, 도면 결과는 기기 수량·연결 관계·HOLD를 순서대로 확인합니다.",
    kind="tip",
)

add_page_break(doc)

# Getting started
add_section_title(doc, "01", "시작 전 설정", "로그인 없이도 계산기와 로컬 BYOK 기능을 사용할 수 있지만, 저장·공유 범위는 기능마다 다릅니다.")
doc.add_heading("1. 접속 및 화면 모드", level=2)
add_steps(
    doc,
    [
        "상단의 시스템 버튼으로 밝은 화면·어두운 화면을 전환합니다.",
        "언어 선택에서 KO/EN을 고릅니다. 계산 단위와 도면 원문은 별도로 확인합니다.",
        "모바일에서는 오른쪽 메뉴 버튼을 눌러 전체 메뉴를 엽니다.",
    ],
)

doc.add_heading("2. BYOK API 키 등록", level=2)
add_steps(
    doc,
    [
        "설정 → BYOK 키 설정으로 이동합니다.",
        "공급자에서 Google을 선택하고 테스트 키를 입력합니다.",
        "모델 목록을 불러온 뒤 도면 분석용 모델을 선택합니다.",
        "기본 호출 호환성 검사를 실행합니다. 이 검사는 도면 판독 정확도를 보증하지 않습니다.",
        "공개 교보재 빠른 분석으로 기기·연결·수치·제안이 실제로 나오는지 확인합니다.",
    ],
)
add_note(
    doc,
    "키 보관",
    "개인 API 키는 브라우저 쪽 암호화 저장을 사용합니다. 공용 PC에서는 작업 후 키를 삭제하고 브라우저 저장 데이터를 정리하십시오. 운영용 키나 회사 공용 키는 관리자 정책 없이 입력하지 마십시오.",
    kind="warn",
)

doc.add_heading("3. 로그인 여부에 따른 차이", level=2)
add_table(
    doc,
    ["기능", "비로그인", "로그인/운영 저장소"],
    [
        ["계산기", "사용 가능", "사용 가능"],
        ["BYOK 분석", "현재 브라우저에서 사용", "정책에 따라 사용"],
        ["최근 작업", "브라우저 세션 범위", "계정·프로젝트 연계 가능"],
        ["보고서 공유", "같은 세션 중심", "운영 저장소가 연결된 경우 지속 보관"],
    ],
    [2200, 3300, 3860],
)

add_page_break(doc)

# Search and answers
add_section_title(doc, "02", "검색과 AI 답변", "질문을 구체적으로 쓰면 계산·기준서·도면 분석으로 자동 분류됩니다.")
doc.add_heading("권장 질문 형식", level=2)
add_table(
    doc,
    ["목적", "입력 예시", "확인할 출력"],
    [
        ["계산", "380V 3상, 50kW, 100m, 동케이블 전압강하 검토", "입력값·공식·단위·판정"],
        ["기준", "KEC 232.3.9 전압강하 조항 원문과 예외", "판본·조항·적용 조건"],
        ["설계 검토", "500kVA 변압기 병렬 운전 조건과 확인 순서", "전제·위험·추가 자료"],
        ["도면", "업로드한 단선도의 모든 기기·관계·미확정 연결을 표로", "수량·관계·HOLD·근거"],
    ],
    [1600, 5000, 2760],
)

doc.add_heading("답변을 검토하는 순서", level=2)
add_steps(
    doc,
    [
        "질문에서 전압·상수·길이·재질·역률·기준 판본이 빠졌는지 봅니다.",
        "AI 설명과 수식 계산 결과가 분리되어 있는지 확인합니다.",
        "출처 조항의 관할·판본·예외 조건을 확인합니다.",
        "‘추정’, ‘가능성’, ‘HOLD’ 문구가 있으면 확정 답으로 사용하지 않습니다.",
        "현장 적용 전 원도면·명판·제조사 자료와 교차 확인합니다.",
    ],
)

add_note(
    doc,
    "좋은 질문의 최소 조건",
    "대상 설비 + 전압/상 + 정격/부하 + 거리/재질 + 적용 기준 + 원하는 결과 형식을 한 문장에 넣으십시오. 예: ‘KEC 2021 기준으로 380V 3상 50kW, 100m 동케이블의 최소 굵기와 전압강하를 표로 검토해 줘.’",
    kind="info",
)

add_page_break(doc)

# Calculators
add_section_title(doc, "03", "계산기", "AI 설명과 별개로 수식 엔진이 결과를 재현합니다.")
add_figure(
    doc,
    TMP / "live-calc-viewport.png",
    "그림 2. 계산기 목록: 분야와 난이도로 57개 계산기를 탐색한다.",
    "전력기초, 전압강하, 케이블, 변압기, 보호협조 등으로 나뉜 계산기 목록 화면.",
    width=6.05,
)

doc.add_heading("계산 절차", level=2)
add_steps(
    doc,
    [
        "계산기 메뉴에서 분야를 선택하거나 검색창에 ‘전압강하’, ‘차단기’처럼 입력합니다.",
        "입력값의 단위와 상 구분을 확인하고 값을 입력합니다.",
        "계산 후 공식·대입값·결과 단위·기준 판정을 함께 확인합니다.",
        "허용치를 넘으면 케이블 상향, 회로 분리, 거리 단축 등 제안을 검토합니다.",
        "계산 이력 또는 프로젝트에 저장한 뒤 원도면과 함께 재검토합니다.",
    ],
)

add_note(
    doc,
    "필수 확인",
    "AI가 도면에서 읽어 온 값이라도 자동으로 확정 입력이 되지는 않습니다. 출처가 한 기기·한 선·한 페이지에 유일하게 결박되지 않으면 계산은 SKIPPED 또는 HOLD로 남습니다.",
    kind="warn",
)

add_page_break(doc)

# SLD overview
add_section_title(doc, "04", "도면 분석: 시작", "이미지·DXF·PDF를 형식에 맞는 분석기로 처리하고, 전체 도면과 구획 판독을 합칩니다.")
add_figure(
    doc,
    TMP / "live-sld.png",
    "그림 3. SLD 도면 분석 화면: 이미지 AI, DXF 벡터, PDF 벡터 분석을 구분한다.",
    "SLD 도면 분석 페이지. 세 가지 분석 탭과 파일 업로드 영역, 전체 문서 판독 버튼이 보인다.",
    width=6.15,
)

doc.add_heading("형식 선택", level=2)
add_table(
    doc,
    ["입력", "권장 경로", "장점", "주의"],
    [
        ["PNG/JPEG", "이미지 AI 분석", "스캔·사진도 처리", "해상도와 모델 품질 영향 큼"],
        ["DXF", "DXF 벡터 파싱", "선·문자 좌표를 직접 추출", "블록·레이어 관례 확인"],
        ["PDF", "PDF 벡터 파싱/전체 분석", "다중 페이지·벡터 문자 활용", "스캔 PDF는 이미지 판독 병행"],
    ],
    [1250, 2400, 2600, 3110],
)

doc.add_heading("업로드 전 체크리스트", level=2)
add_bullets(
    doc,
    [
        "도면 방향이 바르고 여백 때문에 본체가 지나치게 작아지지 않았는가",
        "문자 높이가 확대했을 때 구분되며 P/PPT, B/8 같은 혼동이 없는가",
        "기호·선·접속점이 잘리지 않았고 페이지 번호와 표제란이 보이는가",
        "회사 기밀·개인정보·운영 키가 포함되지 않았는가",
        "분석 목적이 수량, 연결, 계산, 규정 검토 중 무엇인지 정했는가",
    ],
)

add_page_break(doc)

# SLD workflow / zones
add_section_title(doc, "05", "도면 분석: 판독 구조", "전체 이미지 스캔과 정밀 구획 판독을 분리해 누락·중복·경계선 오독을 줄입니다.")
doc.add_heading("분석 파이프라인", level=2)
add_steps(
    doc,
    [
        "전체 페이지를 먼저 훑어 표제란·모선·회로 방향·주요 기기 배치를 잡습니다.",
        "페이지를 논리 구획으로 나누고 각 면에 Pxx-Axx 번호를 부여합니다.",
        "기호 담당, 연결선 담당, 문자 담당, 전기 논리 담당을 서로 다른 호출로 실행합니다.",
        "구획 경계에서 잘린 선은 Pxx-Cxxx, 연결이 확정되지 않은 끝은 Pxx-Uxxx로 표시합니다.",
        "전체 도면을 다시 보며 같은 기기·문자·선을 중복 제거하고 선 조각을 원본 선 ID로 합칩니다.",
        "기기 수량·연결 관계·전압 영역·보호·접지·계산 입력을 교차검증합니다.",
        "근거가 충분한 계산만 실행하고, 부족한 항목은 HOLD와 추가 확인 목록으로 출력합니다.",
    ],
)

add_note(
    doc,
    "구획은 별도 도면이 아닙니다",
    "A 번호는 AI가 정밀 스캔하는 ‘면’, C 번호는 그 면의 경계에 걸린 선, U 번호는 아직 짝을 찾지 못한 끝점입니다. 최종 관계는 구획별 결과를 합친 뒤 전체 도면에서 다시 검증합니다.",
    kind="info",
)

add_figure(
    doc,
    FIG,
    "그림 4. 공개 교보재 단선도 예시: 주 모선·발전기·변압기·리액터·차단기와 전력 흐름 표기를 함께 읽어야 한다.",
    "공개 단선도 예시. 상부 주 모선에 발전기, 3권선 변압기, 리액터와 여러 차단기가 연결되고 MW와 MVAR 값이 표시되어 있다.",
    width=4.55,
)

add_page_break(doc)

# Result interpretation
add_section_title(doc, "06", "도면 결과 읽기", "‘몇 개를 찾았는가’보다 ‘어떤 근거로 무엇과 연결했는가’를 우선 봅니다.")
doc.add_heading("결과 확인 순서", level=2)
add_table(
    doc,
    ["순서", "확인 항목", "질문"],
    [
        ["1", "문서 범위", "모든 페이지와 구획이 처리됐는가?"],
        ["2", "기기 수량", "VCB, TR, BUS, CT/PT, LOAD 수량이 원도면과 맞는가?"],
        ["3", "기기 식별", "Sxx 번호·기기명·정격·페이지 근거가 한 항목에 모였는가?"],
        ["4", "연결 관계", "Lxx가 from→to, 경유 보호기, 방향 근거를 갖는가?"],
        ["5", "경계 연속성", "C 번호가 올바른 짝과 합쳐졌고 U가 남지 않았는가?"],
        ["6", "계산", "전압·전류·길이·규격 등 필수 입력이 원도면에서 확인됐는가?"],
        ["7", "제안", "위험·우선순위·확인 방법이 근거와 함께 제시됐는가?"],
    ],
    [650, 2150, 6560],
)

doc.add_heading("판정 용어", level=2)
add_table(
    doc,
    ["표시", "뜻", "사용자 행동"],
    [
        ["COMPLETE", "요청 범위의 처리와 필수 영수증이 완성됨", "원도면 대조 후 승인 절차 진행"],
        ["PARTIAL", "일부 페이지·구획·역할이 처리됨", "빠진 범위 재실행"],
        ["HOLD", "모호한 기호·문자·선 또는 필수 입력 누락", "추정하지 말고 원도면/명판 확인"],
        ["SKIPPED", "조건 부족으로 계산·검토를 실행하지 않음", "누락 입력 보완"],
        ["FAIL", "파일·파서·호출·계약 위반으로 처리 실패", "오류 원인 수정 후 재실행"],
    ],
    [1300, 4600, 3460],
)

add_note(
    doc,
    "정상적인 HOLD",
    "도면에 값이 없거나 선이 경계에서 여러 갈래로 갈라지는 경우 자동으로 이어 붙이지 않는 것이 맞습니다. HOLD는 분석 실패가 아니라 오확정을 막기 위한 안전 판정일 수 있습니다.",
    kind="tip",
)

add_page_break(doc)

# Suggestions
add_section_title(doc, "07", "분석에서 제안까지", "제안은 ‘무엇을 바꿔라’만 쓰지 않고 근거·영향·추가 자료를 함께 제시해야 합니다.")
doc.add_heading("제안의 4단 구성", level=2)
add_table(
    doc,
    ["구성", "내용", "예시"],
    [
        ["관찰", "도면에서 확인된 사실", "L12 구간 길이 100m, 35㎟ 동케이블"],
        ["판정", "수식·기준·논리 검토", "전압강하 3% 기준 근접"],
        ["제안", "실행 가능한 개선안", "케이블 상향 또는 분전 위치 조정 검토"],
        ["확인", "확정 전 필요한 자료", "부하율·역률·포설조건·판본 확인"],
    ],
    [1400, 3600, 4360],
)

doc.add_heading("제안 품질 체크리스트", level=2)
add_bullets(
    doc,
    [
        "원도면의 Sxx/Lxx/Axx/Cxx/Uxx 근거를 가리키는가",
        "확정 사실과 추정·권고를 문장 수준에서 구분했는가",
        "위험도와 작업 우선순위를 설명하는가",
        "대안별 장단점과 추가 계산 입력을 제시하는가",
        "규정 위반과 단순 권고를 구분하는가",
        "사용자가 다음에 확인할 문서·명판·시험값을 제시하는가",
    ],
)

add_note(
    doc,
    "최종 승인 아님",
    "ESVA의 분석과 제안은 설계 검토를 빠르게 만드는 보조 결과입니다. 보호계전 정정, 차단용량, 접지, 병렬운전처럼 사고 영향이 큰 항목은 원도면·제조사 데이터·시험 성적과 자격 있는 기술자의 검토가 필요합니다.",
    kind="danger",
)

add_page_break(doc)

# Studio
add_section_title(doc, "08", "Studio와 OCR", "명판·문서 이미지를 계산 입력 후보로 바꾸고, 사용자가 확인한 값만 다음 단계로 넘깁니다.")
add_figure(
    doc,
    TMP / "live-studio.png",
    "그림 5. Studio 화면: 이미지·문서에서 정보를 추출한 뒤 계산 또는 검토 작업으로 연결한다.",
    "ESVA Studio 업로드 화면. 문서나 명판 이미지를 올리는 영역과 작업 설명이 보인다.",
    width=6.0,
)

doc.add_heading("권장 사용법", level=2)
add_steps(
    doc,
    [
        "명판이나 문서가 화면에 수평이 되도록 촬영하고 반사광·그림자를 줄입니다.",
        "원본 해상도를 유지한 PNG/JPEG/PDF를 올립니다.",
        "추출된 기기명, 전압, 전류, 용량, 결선, 제조사 표기를 원본과 대조합니다.",
        "P와 PPT, B와 8, 0과 O, kW와 kVA처럼 혼동하기 쉬운 항목을 우선 확인합니다.",
        "확인한 값만 계산기로 보내고, 불명확 값은 HOLD로 남깁니다.",
    ],
)

add_note(
    doc,
    "저해상도 입력",
    "업스케일링은 작은 문자를 보기 쉽게 만들 수 있지만 원본에 없는 획을 복원하지는 못합니다. 한 번에 전체 사진만 맡기기보다 전체 스캔 후 문자·명판 구역을 정밀 분석하고 결과를 다시 합치는 방식이 안전합니다.",
    kind="warn",
)

add_page_break(doc)

# Field and other features
add_section_title(doc, "09", "현장안전·기준서·이력", "작업 조건을 구조화하고, 계산·도면 결과의 근거를 다시 찾을 수 있게 합니다.")
add_figure(
    doc,
    TMP / "live-field.png",
    "그림 6. 현장 안전 체크: 장소·날씨·인원·작업·시간·관리자 정보를 입력한다.",
    "현장 안전 체크 화면. 작업 정보 입력창, 예시 버튼, 안전 분석 시작 버튼이 보인다.",
    width=5.9,
)

doc.add_heading("현장 안전 체크", level=2)
add_steps(
    doc,
    [
        "장소, 날씨, 작업 인원, 작업 종류, 시간, 관리자 수를 한 문장으로 입력합니다.",
        "생성된 산안법/KEC 기반 항목을 작업 전 회의에서 하나씩 확인합니다.",
        "실제 현장 조건과 다른 항목은 수정하고 담당자·확인 시간을 기록합니다.",
        "자동 체크리스트는 현장 위험성평가서와 작업허가서를 대체하지 않습니다.",
    ],
)

doc.add_heading("기준서·용어사전", level=2)
add_bullets(
    doc,
    [
        "기준서: KEC·NEC·IEC 스냅샷을 주제·조항으로 탐색합니다.",
        "다국가 비교: 의무 기준과 권고 기준, 단위 차이를 구분합니다.",
        "용어사전: 한글·영문·약어·동의어를 함께 찾아 도면 표기를 정규화합니다.",
        "중요: 최신 관할 원문 자동 동기화는 아직 보장하지 않으므로 시행일과 개정판을 공식 원문에서 확인합니다.",
    ],
)

doc.add_heading("이력·프로젝트·보고서", level=2)
add_bullets(
    doc,
    [
        "이력: 계산과 검색 기록을 다시 열어 입력값을 비교합니다.",
        "프로젝트: 같은 설비의 계산·도면·보고서를 묶습니다.",
        "보고서: Sxx/Lxx 근거, 수량, 계산, 제안, HOLD를 검토 순서대로 내보냅니다.",
        "비로그인 보고서는 현재 브라우저 세션 중심이므로 다른 기기·새 세션 보관을 전제로 하지 않습니다.",
    ],
)

add_page_break(doc)

# Model guide
add_section_title(doc, "10", "Google 모델 선택 가이드", "아래 표는 공개 교보재 단선도 1종을 모델별 1회 호출한 비교입니다. 일반 정확도나 현장 실증 점수가 아닙니다.")
add_table(
    doc,
    ["모델", "기기/14", "연결/13", "수치/10", "판정"],
    [
        ["gemini-3.5-flash", "14", "13", "9", "성공"],
        ["gemini-3-flash-preview", "14", "13", "10", "성공"],
        ["gemini-3.1-pro-preview-customtools", "14", "13", "10", "성공"],
        ["gemini-3.1-pro-preview", "14", "13", "6", "성공·수치 재확인"],
        ["gemini-3.6-flash", "14", "13", "6", "성공·수치 재확인"],
        ["gemini-flash-latest", "14", "13", "6", "성공·수치 재확인"],
        ["gemini-3.5-flash-lite", "9", "7", "4", "부분 실패"],
        ["gemini-2.5-flash", "3", "0", "2", "실패"],
        ["gemini-flash-lite-latest", "2", "0", "0", "실패"],
        ["gemini-3.1-flash-lite", "0", "0", "0", "실패"],
    ],
    [3900, 1200, 1200, 1200, 1860],
)

doc.add_heading("선택 원칙", level=2)
add_bullets(
    doc,
    [
        "첫 선택: 같은 공개 도면에서 기기·연결을 모두 찾은 상위 모델을 사용합니다.",
        "빠른 확인: Flash 계열로 1차 분석하고 중요한 결과는 상위 모델 또는 다른 공급자로 교차검증합니다.",
        "피할 선택: Lite 모델은 기본 이미지 호출이 성공해도 구조화 SLD 결과가 비거나 관계를 놓칠 수 있습니다.",
        "반복성: 중요한 도면은 같은 모델로 최소 2회 비교하고 기기 수량·관계·HOLD 차이를 기록합니다.",
        "비용/속도: 공개 교보재로 먼저 통과시킨 뒤 회사 도면에 적용합니다.",
    ],
)

add_note(
    doc,
    "모델 목록은 변합니다",
    "공급자가 모델을 추가·폐기하거나 접근 권한을 바꿀 수 있습니다. 설정 화면에서 현재 계정으로 호출 가능한 모델을 다시 불러오고, ‘이미지 입력 성공’과 ‘도면 분석 성공’을 별도로 확인하십시오.",
    kind="warn",
)

add_page_break(doc)

# Troubleshooting
add_section_title(doc, "11", "문제 해결", "오류를 숨기지 말고 입력·모델·도면·저장 범위를 차례로 좁힙니다.")
add_table(
    doc,
    ["증상", "가능한 원인", "조치"],
    [
        ["모델 목록이 비어 있음", "키 오류·권한·네트워크·로그인 상태", "키 재입력 → 기본 호환성 검사 → 공급자 콘솔 권한 확인"],
        ["기기 0개/일부만 반환", "Lite 모델·낮은 해상도·잘못된 탭", "상위 모델 선택 → 원본 해상도 확인 → 전체 분석 재실행"],
        ["선 관계가 끊김", "구획 경계·교차선·분기점 모호", "C/U 번호 확인 → 전체 도면에서 짝 검토 → 원본선 대조"],
        ["PPT가 P/B로 오독", "글자 획 부족·압축·반사광", "문자 구역 정밀 스캔 → OCR 후보 비교 → 원문 확정"],
        ["계산이 SKIPPED", "전압·전류·길이·규격 등 필수 입력 부족", "누락 목록 확인 후 원도면/명판에서 보완"],
        ["결과가 새 세션에서 없음", "비로그인 세션 저장", "로그인·운영 저장소 사용 또는 보고서 내보내기"],
        ["401/403", "인증·모델 접근 권한", "로그인/키/프로젝트 권한 확인"],
        ["500/503", "서버·내구 저장소·외부 API 일시 장애", "원본 보존 → 잠시 후 재개 → 관리자 로그 확인"],
    ],
    [2250, 3100, 4010],
)

doc.add_heading("분석 재시도 순서", level=2)
add_steps(
    doc,
    [
        "원본 파일과 해시를 보존합니다.",
        "문제 페이지·구획·S/L/A/C/U 번호를 기록합니다.",
        "같은 입력으로 동일 모델 1회 재실행해 반복 오류인지 확인합니다.",
        "상위 모델 또는 다른 공급자로 교차검증합니다.",
        "수정한 문자·기기명·관계와 전후 계산 결과를 보고서에 남깁니다.",
    ],
)

add_page_break(doc)

# Final checklist
add_section_title(doc, "12", "작업 완료 체크리스트", "보고서를 내보내기 전에 아래 항목을 모두 확인합니다.")
check_rows = [
    ["□", "입력", "원본 파일·페이지 수·도면 번호·개정판이 맞다"],
    ["□", "범위", "모든 페이지와 필요한 구획이 처리됐다"],
    ["□", "수량", "주요 기기 종류별 수량을 원도면과 대조했다"],
    ["□", "문자", "기기 태그·정격·단위·P/PPT 등 혼동 항목을 확인했다"],
    ["□", "관계", "Sxx/Lxx와 C/U 경계 관계를 확인했다"],
    ["□", "계산", "필수 입력 출처와 단위가 명확하다"],
    ["□", "판정", "COMPLETE/PARTIAL/HOLD/SKIPPED/FAIL을 구분했다"],
    ["□", "제안", "근거·우선순위·대안·추가 자료가 포함됐다"],
    ["□", "기준", "관할·판본·예외 조건을 공식 원문과 대조했다"],
    ["□", "승인", "자격 있는 검토자와 현장 책임자의 확인을 받았다"],
]
add_table(doc, ["확인", "분야", "완료 조건"], check_rows, [900, 1500, 6960])

doc.add_heading("현재 제품 한계", level=2)
add_bullets(
    doc,
    [
        "공개 교보재와 합성 도면에서 코드 경로와 일부 모델 품질을 확인했지만, 독립 정답 라벨을 갖춘 다종 도면의 95% 정확도는 아직 입증되지 않았습니다.",
        "실제 운영 DB·결제·회사 도면·공급자 전 모델을 이용한 운영 실증은 별도입니다.",
        "기준서 최신 원문 자동 동기화, 이메일·푸시 발송, 전역 분산 레이트 리밋은 부분 구현 또는 미연결 상태입니다.",
        "3·4방향 경계 junction은 자동 오병합을 피하기 위해 HOLD로 남을 수 있습니다.",
    ],
)

add_note(
    doc,
    "권장 도입 방식",
    "공개 교보재 → 사내 비기밀 예제 → 독립 검토자가 만든 정답표 → 제한된 실제 업무 순서로 범위를 넓히십시오. 각 단계에서 성공·실패·보류와 수정 내역을 기록하면 회사별 심볼·표기 관례를 안전하게 보정할 수 있습니다.",
    kind="tip",
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("매뉴얼 끝")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor.from_string(MUTED)
set_run_font(r)

# Document properties
props = doc.core_properties
props.title = "ESVA 공개 교보재 기반 도면 분석 사용자 매뉴얼 v0.2.2"
props.subject = "교육용 단선도부터 83쪽 공공기관 전기설계 도면까지의 검색, 계산, 분석, 제안 사용 가이드"
props.author = "ESVA"
props.keywords = "ESVA, 전기, 계산기, 단선도, SLD, 도면 분석, BYOK"
props.comments = "Wikimedia 공개 도면, 공공기관 공개 설계 PDF, 수변전/기능사 교재, 합성 DXF L1~L3 기준"

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
